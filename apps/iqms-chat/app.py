"""
IQMS Chat — Web interface for querying IQMS ERP via Claude + MCP.
Users log in, ask questions in plain English, get answers from Claude
which queries the IQMS Oracle database through the MCP server.
"""

import csv
import io
import json
import logging
import os
import re
import subprocess
import tempfile
import uuid
from collections import deque
from datetime import datetime
from hashlib import sha256
from pathlib import Path
from functools import wraps

from dotenv import load_dotenv

# Load .env from the app directory before anything else reads os.environ.
load_dotenv(Path(__file__).parent / ".env")

import jwt as _jwt

from flask import (
    Flask, render_template, request, redirect, url_for,
    session, flash, jsonify, send_from_directory,
)

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", os.urandom(32))
app.config["MAX_CONTENT_LENGTH"] = 20 * 1024 * 1024  # 20 MB

# ---------------------------------------------------------------------------
# In-memory log buffer (last 500 entries)
# ---------------------------------------------------------------------------
LOG_BUFFER: deque[dict] = deque(maxlen=500)


def _log(level: str, message: str, **extra):
    entry = {
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "level": level,
        "message": message,
        **extra,
    }
    LOG_BUFFER.append(entry)
    # Also print to stdout for systemd journal
    print(f"[{entry['timestamp']}] {level}: {message}", flush=True)


def log_info(message: str, **extra):
    _log("INFO", message, **extra)


def log_error(message: str, **extra):
    _log("ERROR", message, **extra)


def log_warn(message: str, **extra):
    _log("WARN", message, **extra)

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
DATA_DIR = Path(__file__).parent / "data"
DATA_DIR.mkdir(exist_ok=True)
USERS_FILE = DATA_DIR / "users.json"
REPORTS_DIR = Path(__file__).parent / "reports"
REPORTS_DIR.mkdir(exist_ok=True)
UPLOADS_DIR = Path(__file__).parent / "uploads"
UPLOADS_DIR.mkdir(exist_ok=True)

MAX_UPLOAD_SIZE = 20 * 1024 * 1024  # 20 MB
ALLOWED_EXTENSIONS = {
    ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp",
    ".docx", ".xlsx", ".xls", ".csv",
    ".txt", ".md", ".json", ".xml", ".yaml", ".yml",
    ".pdf",
}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}

# Demo mode: the production version of this app spawns Claude (`claude -p`)
# with MCP servers that talk to an IQMS Oracle ERP and a SQL Server SCADA
# database. None of that infrastructure exists in the demo, so /ask returns
# canned responses keyed off keywords — see the bottom of this file.
DEMO_MODE = os.environ.get("DEMO_MODE", "true").lower() == "true"

CLAUDE_MODEL = os.environ.get("CLAUDE_MODEL", "opus")

# ePlant definitions (generic in the demo — original production version had
# three real manufacturing sites)
EPLANTS = {
    "1": {"name": "Plant A", "company": "Acme Industries — Plant A"},
    "2": {"name": "Plant B", "company": "Acme Industries — Plant B"},
    "3": {"name": "Plant C", "company": "Acme Industries — Plant C"},
}

def _load_core_memory() -> str:
    """Stubbed in demo mode — production reads ERP agent memory files."""
    return ""


# Store conversation histories in memory (keyed by session chat_id)
conversations: dict[str, list] = {}

# ---------------------------------------------------------------------------
# User store
# ---------------------------------------------------------------------------

def _load_users() -> dict:
    if USERS_FILE.exists():
        return json.loads(USERS_FILE.read_text())
    return {}


def _save_users(users: dict):
    USERS_FILE.write_text(json.dumps(users, indent=2))


def _hash_pw(password: str, salt: str = "") -> str:
    return sha256(f"{salt}:{password}".encode()).hexdigest()


# ---------------------------------------------------------------------------
# Auth
# ---------------------------------------------------------------------------

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "username" not in session:
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return decorated


def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "username" not in session:
            return redirect(url_for("login"))
        users = _load_users()
        user = users.get(session["username"], {})
        if not user.get("is_admin"):
            flash("Admin access required.", "error")
            return redirect(url_for("chat"))
        return f(*args, **kwargs)
    return decorated


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username", "").strip().lower()
        password = request.form.get("password", "")

        users = _load_users()
        if username not in users:
            flash("Invalid credentials.", "error")
            return render_template("login.html")

        if users[username]["hash"] != _hash_pw(password, users[username]["salt"]):
            flash("Invalid credentials.", "error")
            return render_template("login.html")

        session["username"] = username
        session["display_name"] = users[username].get("display_name", username)
        session["is_admin"] = users[username].get("is_admin", False)
        session["chat_id"] = str(uuid.uuid4())
        session["eplant_id"] = "2"  # Default to Acme
        log_info(f"User '{username}' logged in")
        return redirect(url_for("chat"))

    return render_template("login.html")


@app.route("/sso")
def sso():
    """
    Acme Portal SSO landing. Verifies the portal-issued JWT, then logs the
    visitor in as the hardcoded `admin` user (single-user setup for now;
    real per-user accounts will replace this once IQMS chat gets its own
    user store wired to the directory).
    """
    ptoken = request.args.get("ptoken", "")
    next_path = request.args.get("next", "/")
    if not ptoken:
        flash("Missing SSO token. Open IQMS Chat from the portal.", "error")
        return redirect(url_for("login"))

    secret = os.environ.get("PORTAL_SSO_SECRET", "")
    if not secret:
        log_error("PORTAL_SSO_SECRET not configured — refusing SSO sign-in")
        flash("SSO not configured on this server.", "error")
        return redirect(url_for("login"))

    try:
        claims = _jwt.decode(
            ptoken,
            secret,
            algorithms=["HS256"],
            issuer="acme-portal",
            audience="iqms_chat",
        )
    except _jwt.PyJWTError as exc:
        log_warn(f"Rejected portal SSO token: {exc}")
        flash("Invalid or expired SSO token. Click the IQMS Chat tile in the portal again.", "error")
        return redirect(url_for("login"))

    users = _load_users()
    if "admin" not in users:
        log_error("admin user missing from users.json — cannot complete SSO")
        flash("Admin account not configured.", "error")
        return redirect(url_for("login"))

    admin_user = users["admin"]
    session["username"] = "admin"
    session["display_name"] = claims.get("full_name") or admin_user.get("display_name", "admin")
    session["is_admin"] = bool(admin_user.get("is_admin"))
    session["chat_id"] = str(uuid.uuid4())
    session["eplant_id"] = "2"  # Default to Acme
    session["sso_email"] = claims.get("email", "")
    log_info(f"Portal SSO sign-in for {claims.get('email','?')} → admin")

    if not next_path.startswith("/"):
        next_path = "/"
    return redirect(next_path)


@app.route("/logout")
def logout():
    chat_id = session.get("chat_id")
    if chat_id and chat_id in conversations:
        del conversations[chat_id]
    session.clear()
    return redirect(url_for("login"))


# ---------------------------------------------------------------------------
# ePlant selection
# ---------------------------------------------------------------------------

@app.route("/set-eplant", methods=["POST"])
@login_required
def set_eplant():
    data = request.get_json()
    eplant_id = data.get("eplant_id", "1")
    if eplant_id in EPLANTS:
        session["eplant_id"] = eplant_id
        return jsonify({"ok": True, "name": EPLANTS[eplant_id]["name"]})
    return jsonify({"error": "Invalid ePlant"}), 400


# ---------------------------------------------------------------------------
# File upload & processing
# ---------------------------------------------------------------------------

def _extract_text_from_docx(filepath: Path) -> str:
    """Extract text content from a .docx file."""
    from docx import Document
    doc = Document(str(filepath))
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_text_from_xlsx(filepath: Path) -> str:
    """Extract content from an .xlsx file as CSV-like text."""
    from openpyxl import load_workbook
    wb = load_workbook(str(filepath), read_only=True, data_only=True)
    output = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        output.append(f"=== Sheet: {sheet_name} ===")
        buf = io.StringIO()
        writer = csv.writer(buf)
        for row in ws.iter_rows(values_only=True):
            writer.writerow([str(c) if c is not None else "" for c in row])
        output.append(buf.getvalue())
    wb.close()
    return "\n".join(output)


def _process_upload(filepath: Path) -> dict:
    """Process an uploaded file. Returns dict with 'type' and content info."""
    ext = filepath.suffix.lower()

    if ext in IMAGE_EXTENSIONS:
        # Images: Claude CLI can read them via the Read tool
        return {"type": "image", "path": str(filepath), "name": filepath.name}

    if ext == ".pdf":
        # PDFs: Claude CLI can read them via the Read tool
        return {"type": "pdf", "path": str(filepath), "name": filepath.name}

    if ext == ".docx":
        text = _extract_text_from_docx(filepath)
        return {"type": "text", "content": text[:100000], "name": filepath.name}

    if ext in (".xlsx", ".xls"):
        text = _extract_text_from_xlsx(filepath)
        return {"type": "text", "content": text[:100000], "name": filepath.name}

    # Plain text / csv / json / etc — read directly
    try:
        text = filepath.read_text(encoding="utf-8", errors="replace")
        return {"type": "text", "content": text[:100000], "name": filepath.name}
    except Exception:
        return {"type": "error", "name": filepath.name}


@app.route("/upload", methods=["POST"])
@login_required
def upload():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "No file selected"}), 400

    ext = Path(f.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        return jsonify({"error": f"File type '{ext}' not supported"}), 400

    # Save to user-specific uploads directory
    username = session["username"]
    user_dir = UPLOADS_DIR / username
    user_dir.mkdir(exist_ok=True)

    # Unique filename to avoid collisions
    safe_name = re.sub(r'[^\w\-.]', '_', f.filename)
    unique_name = f"{uuid.uuid4().hex[:8]}_{safe_name}"
    filepath = user_dir / unique_name
    f.save(str(filepath))

    # Check file size
    if filepath.stat().st_size > MAX_UPLOAD_SIZE:
        filepath.unlink()
        return jsonify({"error": "File too large (max 20 MB)"}), 400

    # Process the file
    result = _process_upload(filepath)
    result["upload_id"] = unique_name
    result["original_name"] = f.filename

    log_info(f"File uploaded: {f.filename} ({ext})", user=username, type=result.get("type"))
    return jsonify(result)


# ---------------------------------------------------------------------------
# Chat
# ---------------------------------------------------------------------------

@app.route("/")
@login_required
def chat():
    chat_id = session.get("chat_id", "")
    messages = conversations.get(chat_id, [])
    eplant_id = session.get("eplant_id", "1")
    return render_template("chat.html",
                           username=session["username"],
                           display_name=session.get("display_name", session["username"]),
                           messages=messages,
                           eplants=EPLANTS,
                           current_eplant=eplant_id)


@app.route("/new-chat", methods=["POST"])
@login_required
def new_chat():
    old_id = session.get("chat_id")
    if old_id and old_id in conversations:
        del conversations[old_id]
    session["chat_id"] = str(uuid.uuid4())
    return redirect(url_for("chat"))


# Report marker pattern
REPORT_PATTERN = re.compile(
    r'===REPORT_START===\s*\n(.+?)\n===REPORT_CONTENT===\s*\n(.*?)\n===REPORT_END===',
    re.DOTALL,
)


def _extract_report(answer: str, username: str) -> tuple[str, str | None]:
    """Extract report from answer if present. Returns (cleaned_answer, report_url or None)."""
    match = REPORT_PATTERN.search(answer)
    if not match:
        return answer, None

    filename = match.group(1).strip()
    content = match.group(2).strip()

    # Sanitize filename
    filename = re.sub(r'[^\w\-.]', '_', filename)
    if not filename:
        filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"

    # Save to user-specific reports directory
    user_dir = REPORTS_DIR / username
    user_dir.mkdir(exist_ok=True)
    report_path = user_dir / filename
    report_path.write_text(content, encoding="utf-8")

    # Remove the report markers from the displayed answer
    clean_answer = answer[:match.start()].rstrip()
    report_url = url_for("download_report", filename=filename)

    return clean_answer, report_url


@app.route("/ask", methods=["POST"])
@login_required
def ask():
    data = request.get_json()
    question = data.get("question", "").strip()
    attachments = data.get("attachments", [])
    if not question:
        return jsonify({"error": "Empty question"}), 400

    chat_id = session.get("chat_id", str(uuid.uuid4()))
    if chat_id not in conversations:
        conversations[chat_id] = []

    username = session["username"]

    # Build display content for user message (show attachment names)
    display_content = question
    if attachments:
        file_names = [a.get("original_name", a.get("name", "file")) for a in attachments]
        display_content = question + "\n\n📎 " + ", ".join(file_names)

    # Add user message
    conversations[chat_id].append({
        "role": "user",
        "content": display_content,
        "timestamp": datetime.now().strftime("%H:%M"),
    })

    # Build the full prompt with file contents
    full_prompt = question

    # Collect image/pdf file paths for Claude to read
    file_paths_for_claude = []

    for att in attachments:
        att_type = att.get("type")
        name = att.get("original_name", att.get("name", "file"))

        if att_type == "text":
            # Inline text content
            content = att.get("content", "")
            full_prompt += f"\n\n--- Attached file: {name} ---\n{content}\n--- End of {name} ---"

        elif att_type in ("image", "pdf"):
            # Claude CLI will read these via the Read tool
            fpath = att.get("path", "")
            if fpath and Path(fpath).exists():
                file_paths_for_claude.append(fpath)
                full_prompt += f"\n\n[Attached {att_type}: {name} — saved at {fpath}. Use the Read tool to view it.]"

    eplant_id = session.get("eplant_id", "1")
    eplant = EPLANTS.get(eplant_id, EPLANTS["1"])

    log_info(
        f"Query from '{username}': {question[:120]}{'...' if len(question) > 120 else ''}",
        user=username, eplant=eplant["name"], model=CLAUDE_MODEL,
        attachments=len(attachments), mcp="demo",
    )

    answer = _demo_answer(question, eplant)

    conversations[chat_id].append({
        "role": "assistant",
        "content": answer,
        "timestamp": datetime.now().strftime("%H:%M"),
    })
    return jsonify({"answer": answer})


# ---------------------------------------------------------------------------
# Demo response generator (replaces the production Claude + MCP pipeline)
# ---------------------------------------------------------------------------
DEMO_BANNER = (
    "_(Demo response — the production version of this app sends your "
    "question to Claude with MCP tools that query an Oracle ERP. "
    "This canned answer is illustrative only.)_\n\n"
)

DEMO_RESPONSES: list[tuple[tuple[str, ...], str]] = [
    (("production", "yield", "throughput"),
     "Yesterday's production for **{plant_name}** hit **47,820 lbs** across "
     "12 work orders, vs a daily target of 45,000 lbs (+6.3%). The biggest "
     "contributor was line L-04 (Polymer Resin A), which ran 11.5 hours of "
     "uptime at 92% efficiency.\n\nTop work orders by output:\n"
     "| WO        | Item      | Qty (lbs) | Status     |\n"
     "|-----------|-----------|-----------|------------|\n"
     "| WO-104221 | RES-A-200 | 12,450    | Completed  |\n"
     "| WO-104228 | CAT-B-12  |  9,810    | Completed  |\n"
     "| WO-104231 | RES-A-200 |  8,300    | In progress |\n"),
    (("lot", "qc", "test"),
     "**Plant_A** lot **L-2026-0418-A** is currently in `in_progress` status "
     "with 4 of 5 spec parameters passing:\n\n"
     "- Viscosity:   **3,420 cP**  (spec 3,200 – 3,600) ✓\n"
     "- Moisture:    **0.18 %**     (spec ≤ 0.25) ✓\n"
     "- Density:     **1.142 g/cc** (spec 1.13 – 1.16) ✓\n"
     "- Particle pH: **7.4**         (spec 6.8 – 7.6) ✓\n"
     "- Color (L*):  **84.2**        (spec ≥ 86)  ✗ — out of spec\n\n"
     "Recommend re-test before release."),
    (("shipment", "ship", "freight", "carrier"),
     "Last 7 days at **{plant_name}**: **23 shipments**, total billed "
     "freight **$48,615**, avg $2,114 per shipment. Top lanes:\n\n"
     "| Lane                      | Shipments | Total Cost |\n"
     "|---------------------------|-----------|------------|\n"
     "| Plant A → Atlanta, GA     |  6        |  $13,820   |\n"
     "| Plant A → Cleveland, OH   |  4        |   $7,440   |\n"
     "| Plant A → Houston, TX     |  3        |   $9,210   |\n\n"
     "Carrier mix: FedEx Demo 48%, R+L Demo 32%, UPS Demo 20%."),
    (("inventory", "stock", "on hand"),
     "Current finished-goods inventory snapshot at **{plant_name}**:\n\n"
     "- Polymer Resin A:  **62,400 lbs** (≈ 7.2 days of cover)\n"
     "- Catalyst B-12:    **18,910 lbs** (≈ 4.1 days)\n"
     "- Solvent C:        **31,200 lbs** (≈ 9.5 days)\n\n"
     "Two items below 5-day cover threshold — consider rebuilding."),
    (("customer", "top", "sales"),
     "Top 5 customers YTD at **{plant_name}** by revenue:\n\n"
     "1. Northwind Corp     —  $1.42 M  (28 orders)\n"
     "2. Vandelay Plastics  —  $1.18 M  (22 orders)\n"
     "3. Globex Industries  —    $872 K (19 orders)\n"
     "4. Hooli Materials    —    $610 K (14 orders)\n"
     "5. Initech Polymers   —    $540 K (12 orders)\n"),
]

DEMO_FALLBACK = (
    "I'm running in **demo mode**, so I can't actually query an ERP. "
    "Try asking about one of these to see a sample answer:\n\n"
    "- *what was production yesterday at {plant_name}?*\n"
    "- *show me lot L-2026-0418-A test results*\n"
    "- *which shipments went out last week?*\n"
    "- *what's our inventory of Polymer Resin A?*\n"
    "- *who are our top customers this year?*\n"
)


def _demo_answer(question: str, eplant: dict) -> str:
    q = question.lower()
    plant_name = eplant.get("name", "Plant A")
    for keywords, body in DEMO_RESPONSES:
        if any(kw in q for kw in keywords):
            return DEMO_BANNER + body.format(plant_name=plant_name)
    return DEMO_BANNER + DEMO_FALLBACK.format(plant_name=plant_name)


# ---------------------------------------------------------------------------
# Logs (admin only)
# ---------------------------------------------------------------------------

@app.route("/logs")
@admin_required
def view_logs():
    return render_template("logs.html",
                           username=session["username"],
                           display_name=session.get("display_name", session["username"]))


@app.route("/api/logs")
@admin_required
def api_logs():
    level = request.args.get("level", "")
    user = request.args.get("user", "")
    limit = int(request.args.get("limit", 200))

    entries = list(LOG_BUFFER)
    if level:
        entries = [e for e in entries if e["level"] == level.upper()]
    if user:
        entries = [e for e in entries if e.get("user", "").lower() == user.lower()]
    entries = entries[-limit:]
    return jsonify(entries)


# ---------------------------------------------------------------------------
# Report downloads
# ---------------------------------------------------------------------------

@app.route("/reports/<filename>")
@login_required
def download_report(filename):
    username = session["username"]
    user_dir = REPORTS_DIR / username
    if not (user_dir / filename).exists():
        return "Report not found.", 404
    return send_from_directory(user_dir, filename, as_attachment=True)


# ---------------------------------------------------------------------------
# Export — convert markdown response to Excel / Word / PDF
# ---------------------------------------------------------------------------

def _md_to_html(md_text: str) -> str:
    """Convert markdown to styled HTML for PDF/Word rendering."""
    import markdown
    body = markdown.markdown(md_text, extensions=["tables", "fenced_code"])
    return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
  body {{ font-family: 'Segoe UI', Arial, sans-serif; font-size: 11pt; color: #1e293b;
         max-width: 800px; margin: 2rem auto; padding: 0 1rem; }}
  h1,h2,h3 {{ color: #92400e; border-bottom: 1px solid #e5e7eb; padding-bottom: 0.3rem; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1rem 0; font-size: 10pt; }}
  th {{ background: #f59e0b; color: #000; font-weight: 600; }}
  th, td {{ border: 1px solid #d1d5db; padding: 6px 10px; text-align: left; }}
  tr:nth-child(even) {{ background: #f9fafb; }}
  code {{ background: #f1f5f9; padding: 2px 5px; border-radius: 3px; font-size: 0.9em; }}
  pre {{ background: #f1f5f9; padding: 1rem; border-radius: 6px; overflow-x: auto; font-size: 0.85em; }}
  strong {{ color: #92400e; }}
</style></head><body>{body}</body></html>"""


def _parse_tables_from_md(md_text: str) -> list[list[list[str]]]:
    """Extract markdown tables as lists of rows. Each table is a list of rows, each row a list of cells."""
    tables = []
    current_table = []
    for line in md_text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            # Skip separator rows (|---|---|)
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if cells and not all(set(c) <= {"-", ":", " "} for c in cells):
                current_table.append(cells)
        else:
            if current_table:
                tables.append(current_table)
                current_table = []
    if current_table:
        tables.append(current_table)
    return tables


@app.route("/export", methods=["POST"])
@login_required
def export():
    data = request.get_json()
    content = data.get("content", "")
    fmt = data.get("format", "pdf")
    if not content:
        return jsonify({"error": "No content to export"}), 400

    username = session["username"]
    user_dir = REPORTS_DIR / username
    user_dir.mkdir(exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")

    try:
        if fmt == "pdf":
            import weasyprint
            html = _md_to_html(content)
            filename = f"export_{ts}.pdf"
            filepath = user_dir / filename
            weasyprint.HTML(string=html).write_pdf(str(filepath))

        elif fmt == "docx":
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.table import WD_TABLE_ALIGNMENT

            doc = Document()
            style = doc.styles["Normal"]
            style.font.name = "Calibri"
            style.font.size = Pt(11)

            tables = _parse_tables_from_md(content)
            table_idx = 0

            for line in content.split("\n"):
                stripped = line.strip()

                # Skip table rows (handled separately)
                if stripped.startswith("|") and stripped.endswith("|"):
                    cells = [c.strip() for c in stripped.strip("|").split("|")]
                    if all(set(c) <= {"-", ":", " "} for c in cells):
                        continue
                    # Check if this is the first row of a table we haven't rendered yet
                    if table_idx < len(tables) and cells == tables[table_idx][0]:
                        tbl_data = tables[table_idx]
                        table_idx += 1
                        tbl = doc.add_table(rows=len(tbl_data), cols=len(tbl_data[0]))
                        tbl.style = "Table Grid"
                        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                        for i, row_data in enumerate(tbl_data):
                            for j, cell_text in enumerate(row_data):
                                if j < len(tbl.columns):
                                    cell = tbl.rows[i].cells[j]
                                    cell.text = cell_text
                                    if i == 0:
                                        for run in cell.paragraphs[0].runs:
                                            run.bold = True
                        doc.add_paragraph()
                    continue

                if stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    doc.add_paragraph(stripped[2:], style="List Bullet")
                elif stripped:
                    # Handle bold markers
                    p = doc.add_paragraph()
                    parts = stripped.split("**")
                    for k, part in enumerate(parts):
                        if part:
                            run = p.add_run(part)
                            if k % 2 == 1:
                                run.bold = True

            filename = f"export_{ts}.docx"
            filepath = user_dir / filename
            doc.save(str(filepath))

        elif fmt == "xlsx":
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

            wb = Workbook()
            ws = wb.active
            ws.title = "Export"

            header_font = Font(name="Calibri", bold=True, color="000000", size=11)
            header_fill = PatternFill(start_color="F59E0B", end_color="F59E0B", fill_type="solid")
            thin_border = Border(
                left=Side(style="thin"), right=Side(style="thin"),
                top=Side(style="thin"), bottom=Side(style="thin"),
            )

            tables = _parse_tables_from_md(content)

            if tables:
                row_num = 1
                for t_idx, tbl in enumerate(tables):
                    if t_idx > 0:
                        row_num += 1  # blank row between tables

                    for i, row_data in enumerate(tbl):
                        for j, cell_text in enumerate(row_data):
                            cell = ws.cell(row=row_num, column=j + 1, value=cell_text)
                            cell.border = thin_border
                            cell.alignment = Alignment(wrap_text=True)
                            if i == 0:
                                cell.font = header_font
                                cell.fill = header_fill
                            # Try to convert numbers
                            try:
                                cell.value = float(cell_text.replace(",", "").replace("$", ""))
                                cell.number_format = '#,##0.00' if "." in cell_text else '#,##0'
                            except (ValueError, AttributeError):
                                pass
                        row_num += 1

                # Auto-fit column widths
                for col_cells in ws.columns:
                    max_len = max((len(str(c.value or "")) for c in col_cells), default=10)
                    ws.column_dimensions[col_cells[0].column_letter].width = min(max_len + 3, 50)
            else:
                # No tables found — dump text line by line
                for i, line in enumerate(content.split("\n"), 1):
                    ws.cell(row=i, column=1, value=line.strip())

            filename = f"export_{ts}.xlsx"
            filepath = user_dir / filename
            wb.save(str(filepath))

        else:
            return jsonify({"error": f"Unknown format: {fmt}"}), 400

        log_info(f"Export {fmt.upper()} for '{username}'", filename=filename)
        return jsonify({"url": url_for("download_report", filename=filename)})

    except Exception as e:
        log_error(f"Export failed: {str(e)}", user=username, format=fmt)
        return jsonify({"error": f"Export failed: {str(e)}"}), 500


# ---------------------------------------------------------------------------
# Admin — User Management
# ---------------------------------------------------------------------------

@app.route("/admin")
@admin_required
def admin():
    users = _load_users()
    return render_template("admin.html",
                           username=session["username"],
                           display_name=session.get("display_name", session["username"]),
                           users=users)


@app.route("/admin/add-user", methods=["POST"])
@admin_required
def add_user():
    username = request.form.get("username", "").strip().lower()
    display_name = request.form.get("display_name", "").strip()
    password = request.form.get("password", "")
    is_admin = request.form.get("is_admin") == "on"

    if not username or not password:
        flash("Username and password are required.", "error")
        return redirect(url_for("admin"))

    if len(password) < 4:
        flash("Password must be at least 4 characters.", "error")
        return redirect(url_for("admin"))

    users = _load_users()
    if username in users:
        flash(f"User '{username}' already exists.", "error")
        return redirect(url_for("admin"))

    salt = os.urandom(16).hex()
    users[username] = {
        "hash": _hash_pw(password, salt),
        "salt": salt,
        "display_name": display_name or username,
        "is_admin": is_admin,
        "created": datetime.now().isoformat(),
    }
    _save_users(users)
    flash(f"User '{display_name or username}' created.", "success")
    return redirect(url_for("admin"))


@app.route("/admin/delete-user", methods=["POST"])
@admin_required
def delete_user():
    username = request.form.get("username", "")
    if username == session["username"]:
        flash("You can't delete yourself.", "error")
        return redirect(url_for("admin"))

    users = _load_users()
    if username in users:
        name = users[username].get("display_name", username)
        del users[username]
        _save_users(users)
        flash(f"User '{name}' deleted.", "success")
    return redirect(url_for("admin"))


@app.route("/admin/reset-password", methods=["POST"])
@admin_required
def reset_password():
    username = request.form.get("username", "")
    new_password = request.form.get("new_password", "")

    if len(new_password) < 4:
        flash("Password must be at least 4 characters.", "error")
        return redirect(url_for("admin"))

    users = _load_users()
    if username not in users:
        flash("User not found.", "error")
        return redirect(url_for("admin"))

    salt = os.urandom(16).hex()
    users[username]["hash"] = _hash_pw(new_password, salt)
    users[username]["salt"] = salt
    _save_users(users)
    flash(f"Password reset for '{users[username].get('display_name', username)}'.", "success")
    return redirect(url_for("admin"))


# ---------------------------------------------------------------------------
# Bootstrap admin if no users exist
# ---------------------------------------------------------------------------
def _ensure_admin():
    users = _load_users()
    if not users:
        salt = os.urandom(16).hex()
        users["admin"] = {
            "hash": _hash_pw("admin", salt),
            "salt": salt,
            "display_name": "Administrator",
            "is_admin": True,
            "created": datetime.now().isoformat(),
        }
        _save_users(users)
        print(">>> Default admin account created (username: admin, password: admin)")
        print(">>> Change the password immediately via the admin panel!")


# ---------------------------------------------------------------------------
# Run
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    _ensure_admin()
    app.run(host="0.0.0.0", port=5055, debug=False)
